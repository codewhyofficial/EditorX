from flask import Flask, request, jsonify
from werkzeug.utils import secure_filename
from docx import Document
from flask_cors import CORS  # Import CORS
import os
from docx.enum.text import WD_UNDERLINE
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.shape import WD_INLINE_SHAPE
# from docx.enum.text import WD_LINE_STYLE

app = Flask(__name__)
# Initialize CORS and specifically allow localhost:5173
CORS(app, resources={r"/*": {"origins": ["http://localhost:5173"]}})

# Step 1: Create the Upload Folder
upload_folder_path = './upload'
if not os.path.exists(upload_folder_path):
    os.makedirs(upload_folder_path)  # This creates the folder including any necessary parent directories

app.config['UPLOAD_FOLDER'] = upload_folder_path  # Update with your upload folder path
app.config['ALLOWED_EXTENSIONS'] = {'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
        return jsonify({'filename': filename}), 200
    else:
        return jsonify({'error': 'File type not allowed, please upload a .docx file'}), 400

# @app.route('/get_document/<filename>', methods=['GET'])
# def get_document(filename):
#     doc_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
#     doc = Document(doc_path)
#     full_text = []
#     for paragraph in doc.paragraphs:
#         full_text.append(paragraph.text)
    
#     return jsonify({'content': '\n'.join(full_text)}), 200

# final 
@app.route('/get_document/<filename>', methods=['GET'])
def get_document(filename):
    doc_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    if os.path.exists(doc_path):
        doc = Document(doc_path)
        content = []

        # Handle paragraphs
        for paragraph in doc.paragraphs:
            para_data = {
                'type': 'paragraph',
                'runs': []
            }
            for run in paragraph.runs:
                run_data = {
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font': {
                        'name': run.font.name,
                        'size': run.font.size.pt if run.font.size else None
                    }
                }
                para_data['runs'].append(run_data)
            content.append(para_data)

        # Handle tables
        for table in doc.tables:
            table_data = {
                'type': 'table',
                'rows': []
            }
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_data = {
                        'paragraphs': []
                    }
                    for paragraph in cell.paragraphs:
                        para_data = {
                            'runs': []
                        }
                        for run in paragraph.runs:
                            run_data = {
                                'text': run.text,
                                'bold': run.bold,
                                'italic': run.italic,
                                'underline': run.underline,
                                'font': {
                                    'name': run.font.name,
                                    'size': run.font.size.pt if run.font.size else None
                                }
                            }
                            para_data['runs'].append(run_data)
                        cell_data['paragraphs'].append(para_data)
                    row_data.append(cell_data)
                table_data['rows'].append(row_data)
            content.append(table_data)

        return jsonify({'content': content}), 200
    return jsonify({'message': 'File not found'}), 404


# @app.route('/get_document/<filename>', methods=['GET'])
# def get_document(filename):
#     doc_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
#     doc = Document(doc_path)
#     content = []

#     # Function to extract run data
#     def extract_run_data(run):
#         run_data = {
#             'text': run.text,
#             'bold': run.bold,
#             'italic': run.italic,
#             'underline': run.underline,
#             'font': {
#                 'name': run.font.name,
#                 'size': run.font.size.pt if run.font.size else None  # Convert to points (pt)
#             }
#         }

#         if run.underline:
#             if run.underline == WD_UNDERLINE.DASH:
#                 run_data['underline_style'] = 'dash'
#             elif run.underline == WD_UNDERLINE.DOTTED:
#                 run_data['underline_style'] = 'dotted'
#             elif run.underline == WD_UNDERLINE.SINGLE:
#                 run_data['underline_style'] = 'single'
#             elif run.underline == WD_UNDERLINE.WAVY:
#                 run_data['underline_style'] = 'wavy'

#         return run_data

#     # Extract paragraphs
#     for paragraph in doc.paragraphs:
#         para_data = {
#             'text': '',
#             'runs': [extract_run_data(run) for run in paragraph.runs]
#         }
#         content.append(para_data)

#     # Extract tables
#     for table in doc.tables:
#         table_data = []
#         for row in table.rows:
#             row_data = []
#             for cell in row.cells:
#                 cell_data = {
#                     'text': cell.text,
#                     'runs': []
#                 }
#                 for paragraph in cell.paragraphs:
#                     for run in paragraph.runs:
#                         cell_data['runs'].append(extract_run_data(run))
#                 row_data.append(cell_data)
#             table_data.append(row_data)
#         content.append({'table': table_data})

#     return jsonify({'content': content}), 200





# @app.route('/get_document/<filename>', methods=['GET'])
# def get_document(filename):
#     doc_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
#     doc = Document(doc_path)
#     content = []

#     for paragraph in doc.paragraphs:
#         para_data = {
#             'text': '',
#             'runs': []
#         }

#         for run in paragraph.runs:
#             run_data = {
#                 'text': run.text,
#                 'bold': run.bold,
#                 'italic': run.italic,
#                 'underline': run.underline,
#                 'font': {
#                     'name': run.font.name,
#                     'size': run.font.size.pt if run.font.size else None  # Convert to points (pt)
#                 }
#             }

#             if run.underline:
#                 if run.underline == WD_UNDERLINE.DASH:
#                     run_data['underline_style'] = 'dash'
#                 elif run.underline == WD_UNDERLINE.DOTTED:
#                     run_data['underline_style'] = 'dotted'
#                 elif run.underline == WD_UNDERLINE.SINGLE:
#                     run_data['underline_style'] = 'single'
#                 elif run.underline == WD_UNDERLINE.WAVY:
#                     run_data['underline_style'] = 'wavy'

#             para_data['runs'].append(run_data)

#         content.append(para_data)

#     for table in doc.tables:
#         table_data = []
#         for row in table.rows:
#             row_data = []
#             for cell in row.cells:
#                 cell_data = {
#                     'text': cell.text,
#                     'runs': []
#                 }
#                 for paragraph in cell.paragraphs:
#                     for run in paragraph.runs:
#                         run_data = {
#                             'text': run.text,
#                             'bold': run.bold,
#                             'italic': run.italic,
#                             'underline': run.underline,
#                             'font': {
#                                 'name': run.font.name,
#                                 'size': run.font.size.pt if run.font.size else None  # Convert to points (pt)
#                             }
#                         }

#                         if run.underline:
#                             if run.underline == WD_UNDERLINE.DASH:
#                                 run_data['underline_style'] = 'dash'
#                             elif run.underline == WD_UNDERLINE.DOTTED:
#                                 run_data['underline_style'] = 'dotted'
#                             elif run.underline == WD_UNDERLINE.SINGLE:
#                                 run_data['underline_style'] = 'single'
#                             elif run.underline == WD_UNDERLINE.WAVY:
#                                 run_data['underline_style'] = 'wavy'

#                         cell_data['runs'].append(run_data)
#                 row_data.append(cell_data)
#             table_data.append(row_data)
#         content.append({'table': table_data})

#     for shape in doc.inline_shapes:
#         shape_data = {
#             'type': shape.type,
#             'width': shape.width.pt,
#             'height': shape.height.pt
#         }
#         if shape.line and shape.line.style:
#             if shape.line.style == WD_LINE_STYLE.SINGLE:
#                 shape_data['line_type'] = 'single'
#             elif shape.line.style == WD_LINE_STYLE.THICK_THIN:
#                 shape_data['line_type'] = 'thick_thin'
#             elif shape.line.style == WD_LINE_STYLE.THIN_THICK:
#                 shape_data['line_type'] = 'thin_thick'
#             elif shape.line.style == WD_LINE_STYLE.TRIPLE:
#                 shape_data['line_type'] = 'triple'
#         content.append({'shape': shape_data})

#     return jsonify({'content': content}), 200




# @app.route('/get_document/<filename>', methods=['GET'])
# def get_docx_content(filename):
#     file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
#     if os.path.exists(file_path):
#         doc = Document(file_path)
#         content = []

#         for para in doc.paragraphs:
#             para_content = {
#                 'text': para.text,
#                 'bold': [run.bold for run in para.runs],
#                 'italic': [run.italic for run in para.runs],
#                 'underline': [run.underline for run in para.runs],
#                 'spacing': para.paragraph_format.space_after.pt if para.paragraph_format.space_after else None,
#                 'border': {
#                     'top': para.paragraph_format.border_top.width if para.paragraph_format.border_top else None,
#                     'bottom': para.paragraph_format.border_bottom.width if para.paragraph_format.border_bottom else None,
#                     'left': para.paragraph_format.border_left.width if para.paragraph_format.border_left else None,
#                     'right': para.paragraph_format.border_right.width if para.paragraph_format.border_right else None
#                 }
#             }
#             content.append(para_content)
#         return jsonify({'content': content})
#     return jsonify({'message': 'File not found'}), 404

# @app.route('/save_document/<filename>', methods=['POST'])
# def save_document(filename):
#     doc_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
#     doc = Document(doc_path)

#     data = request.get_json()
#     new_content = data.get('content')

#     paragraphs = new_content.split('\n')
#     for i, paragraph in enumerate(paragraphs):
#         if i < len(doc.paragraphs):
#             doc.paragraphs[i].text = paragraph
#         else:
#             doc.add_paragraph(paragraph)

#     doc.save(doc_path)
#     return jsonify({'message': 'Document saved successfully'}), 200

@app.route('/save_document/<filename>', methods=['POST'])
def save_document(filename):
    data = request.json
    content = data.get('content', [])
    doc_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    doc = Document()

    for item in content:
        if item['type'] == 'paragraph':
            p = doc.add_paragraph()
            for run in item['runs']:
                r = p.add_run(run['text'])
                r.bold = run['bold']
                r.italic = run['italic']
                r.underline = run['underline']
                if run['font']['name']:
                    r.font.name = run['font']['name']
                if run['font']['size']:
                    r.font.size = run['font']['size']

        elif item['type'] == 'table':
            table = doc.add_table(rows=0, cols=len(item['rows'][0]))
            for row in item['rows']:
                tr = table.add_row()
                for idx, cell in enumerate(row):
                    tc = tr.cells[idx]
                    for paragraph in cell['paragraphs']:
                        p = tc.add_paragraph()
                        for run in paragraph['runs']:
                            r = p.add_run(run['text'])
                            r.bold = run['bold']
                            r.italic = run['italic']
                            r.underline = run['underline']
                            if run['font']['name']:
                                r.font.name = run['font']['name']
                            if run['font']['size']:
                                r.font.size = run['font']['size']

    doc.save(doc_path)
    return jsonify({'message': 'Document saved successfully', 'content': content}), 200

if __name__ == '__main__':
    app.run(debug=True)
